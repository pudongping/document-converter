# -*- coding: utf-8 -*-

'''
文件转换为 docx
'''

__author__ = 'alex'

import logging

from docx import Document
import pdfplumber

from app.helper import *
from app.config.app import AppConfig


class ConverterDocx(object):

    def pdf_to_docx(self, input_path='', output_path=''):
        if (not bool(input_path)) and (not bool(output_path)): return False
        # 查找指定输入路径下所有的 .pdf 文件
        file_list = find_all_file_by_path(AppConfig.PDF, input_path)

        counter = 0
        for file in file_list:
            file_name = os.path.split(file)[1]  # 文件名 eg: demo1.pdf
            docx_name_with_path = output_path + file_name + '.' + AppConfig.DOCX  # 在原文件末尾加上 .docx 后缀
            msg = '正在处理 <%s> 转换后的文件保存路径为： <%s>' % (file, docx_name_with_path)
            print(msg)
            logging.info(msg)

            with pdfplumber.open(file) as pdf:
                print('正在处理文件 <%s>， 一共 %s 页' % (file, len(pdf.pages)))
                content = ''
                flag = True
                for i in range(len(pdf.pages)):
                    print('正在处理文件 <%s> 第 %s 页' % (file, i))
                    page = pdf.pages[i]
                    if page.extract_text() == None:
                        print(f'{file} 是图片拼接起来的，抱歉无法转换！')
                        flag = False
                        break
                    page_content = '\n'.join(page.extract_text().split('\n')[:-1])
                    content = content + page_content
                    if os.path.exists(docx_name_with_path):
                        doc = Document(docx_name_with_path)
                    else:
                        doc = Document()
                    doc.add_paragraph(content)
                    doc.save(docx_name_with_path)
                    content = ''
                    print('文件 <%s> 第 %s 页处理完成！' % (file, i))
                if flag:
                    print('文件 <%s> 处理完成，保存位置为 <%s>' % (file, docx_name_with_path))

            counter += 1

        counter_msg = f'总共需要处理 {len(file_list)} 个文件，{counter} 个文件已经处理完毕！'
        print(counter_msg)
        logging.info(counter_msg)
